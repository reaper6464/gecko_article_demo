# app.py — Gecko Technical Knowledge Base Demo (with Password Gate)
import os
import re
import pandas as pd
import streamlit as st
import csv
from docx import Document

# ---- PASSWORD PROTECTION ----
PASSWORD = "gecko-demo"  # Change to your preferred demo password!

def password_gate():
    if "authenticated" not in st.session_state:
        st.session_state.authenticated = False
    if not st.session_state.authenticated:
        st.title("Gecko Technical Knowledge Base")
        pwd = st.text_input("Enter password:", type="password")
        if st.button("Login"):
            if pwd == PASSWORD:
                st.session_state.authenticated = True
                st.experimental_rerun()
            else:
                st.error("Incorrect password.")
        st.stop()

password_gate()

# ---- MAIN UI STARTS HERE ----

st.set_page_config(layout="wide")
st.sidebar.title("Navigation")
page = st.sidebar.radio("Go to", ["Review Articles"])

# Directory selection
default_save_path = st.sidebar.text_input("Approved articles save path:", value=os.getcwd())

PRODUCT_LIST = [
    "in.ye", "in.xe", "in.touch 2", "in.touch 3", "in.k1000", "Gecko Waterlab", "pumps"
]
CATEGORY_LIST = ["Bug/Troubleshooting", "Upcoming Feature", "Knowledge"]

REJECT_FILE = "rejected_articles.csv"
if not os.path.isfile(REJECT_FILE):
    with open(REJECT_FILE, "w", encoding="utf-8", newline='') as f:
        writer = csv.writer(f)
        writer.writerow([
            "Article Index", "Draft Index", "Title", "Category", "Product", "Body", "Rejection Reason"
        ])

def sanitize_filename(title):
    return re.sub(r'[\\/*?:"<>|]', '_', title)

def parse_article_fields(article_block):
    fields = {"title": "", "category": "", "product": "", "body": ""}
    lines = article_block.splitlines()
    buffer = []
    for line in lines:
        if line.lower().startswith("title:"):
            fields["title"] = line[6:].strip()
        elif line.lower().startswith("category:"):
            fields["category"] = line[9:].strip()
        elif line.lower().startswith("product:"):
            fields["product"] = line[8:].strip()
        else:
            buffer.append(line)
    fields["body"] = "\n".join(buffer).strip()
    return fields

def compose_article_block(fields):
    return f"Title: {fields['title']}\nCategory: {fields['category']}\nProduct: {fields['product']}\n\n{fields['body']}\n"

if page == "Review Articles":
    st.header("Review Extracted Technical Articles")
    # --- LOAD DATA ---
    try:
        draft_df = pd.read_csv("draft_articles.csv")
    except Exception as e:
        st.error(f"Could not load draft_articles.csv: {e}")
        st.stop()

    all_articles = []
    for idx, row in draft_df.iterrows():
        article_blocks = row["articles"].split('---')
        for i, article in enumerate(article_blocks):
            article = article.strip()
            if not article:
                continue
            fields = parse_article_fields(article)
            all_articles.append({
                "df_idx": idx,
                "art_idx": i + 1,
                "email": row["email"],
                "fields": fields
            })

    # ---- FILTERS ----
    selected_product = st.sidebar.selectbox(
        "Filter by product (optional):", ["All"] + PRODUCT_LIST
    )
    selected_category = st.sidebar.selectbox(
        "Filter by category (optional):", ["All"] + CATEGORY_LIST
    )
    search_text = st.sidebar.text_input("Search in title/body:")

    def matches_filter(fields):
        if selected_product != "All" and selected_product not in fields["product"]:
            return False
        if selected_category != "All" and selected_category != fields["category"]:
            return False
        if search_text:
            stext = search_text.lower()
            if stext not in fields["title"].lower() and stext not in fields["body"].lower():
                return False
        return True

    num_shown = sum(matches_filter(art['fields']) for art in all_articles)
    st.info(f"Filtered and showing {num_shown} article(s).")

    for idx, art in enumerate(all_articles):
        fields = art["fields"]
        if not matches_filter(fields):
            continue

        st.markdown("---")
        st.subheader(f"Article {art['df_idx']+1}.{art['art_idx']}")
        with st.expander("Show source email", expanded=False):
            st.text_area(
                "Email",
                value=art["email"],
                height=150,
                disabled=True,
                key=f"email_{art['df_idx']}_{art['art_idx']}"
            )

        # Editable fields
        new_title = st.text_input(
            f"Title {art['df_idx']+1}.{art['art_idx']}", value=fields["title"], key=f"title_{art['df_idx']}_{art['art_idx']}"
        )
        new_category = st.selectbox(
            f"Category {art['df_idx']+1}.{art['art_idx']}", CATEGORY_LIST,
            index=CATEGORY_LIST.index(fields["category"]) if fields["category"] in CATEGORY_LIST else 0,
            key=f"cat_{art['df_idx']}_{art['art_idx']}"
        )
        new_product = st.multiselect(
            f"Product(s) {art['df_idx']+1}.{art['art_idx']}", PRODUCT_LIST,
            default=[p for p in PRODUCT_LIST if p in fields["product"]],
            key=f"prod_{art['df_idx']}_{art['art_idx']}"
        )
        new_body = st.text_area(
            f"Body {art['df_idx']+1}.{art['art_idx']}", value=fields["body"], height=300, key=f"body_{art['df_idx']}_{art['art_idx']}"
        )

        # Approve & Save as docx
        if st.button(f"Approve and Save Article {art['df_idx']+1}.{art['art_idx']}"):
            fields_out = {
                "title": new_title.strip(),
                "category": new_category,
                "product": ", ".join(new_product),
                "body": new_body.strip()
            }
            safe_title = sanitize_filename(fields_out["title"] or f"article_{art['df_idx']+1}_{art['art_idx']}")
            save_dir = default_save_path if os.path.isdir(default_save_path) else os.getcwd()
            doc = Document()
            doc.add_heading(fields_out["title"], level=1)
            doc.add_paragraph(f"Category: {fields_out['category']}")
            doc.add_paragraph(f"Product(s): {fields_out['product']}")
            doc.add_paragraph("")
            for paragraph in fields_out["body"].split("\n"):
                doc.add_paragraph(paragraph)
            docx_path = os.path.join(save_dir, f"approved_{safe_title}.docx")
            doc.save(docx_path)
            st.success(f"Saved as {docx_path}")

        # Rejection state
        if "reject_states" not in st.session_state:
            st.session_state.reject_states = {}

        reject_key = f"rej_{art['df_idx']}_{art['art_idx']}"

        if st.button(f"Reject Article {art['df_idx']+1}.{art['art_idx']}"):
            st.session_state.reject_states[reject_key] = True

        if st.session_state.reject_states.get(reject_key, False):
            rejection_reason = st.text_input(
                f"Reason for rejecting Article {art['df_idx']+1}.{art['art_idx']} (required):",
                key=f"rej_reason_{art['df_idx']}_{art['art_idx']}"
            )
            if st.button(f"Confirm Rejection {art['df_idx']+1}.{art['art_idx']}"):
                if not rejection_reason.strip():
                    st.error("You must provide a reason for rejection.")
                else:
                    with open(REJECT_FILE, "a", encoding="utf-8", newline='') as f:
                        writer = csv.writer(f)
                        writer.writerow([
                            art['df_idx']+1,
                            art['art_idx'],
                            fields["title"],
                            fields["category"],
                            fields["product"],
                            fields["body"],
                            rejection_reason.strip()
                        ])
                    st.warning("Article rejected and reason saved.")
                    st.session_state.reject_states[reject_key] = False
